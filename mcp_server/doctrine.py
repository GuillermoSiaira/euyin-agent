"""
doctrine.py — Recuperación doctrinal sobre el corpus completo del repo ai-oracle.

Indexa los textos doctrinales/conceptuales (Axiomática, system prompt de Lilly,
técnicas persas, Grimoire, el Canon en .docx, formalización, diálogos de
paradigma) y devuelve los fragmentos más relevantes para un tema/tags por
solapamiento de palabras clave.

Retrieval léxico, sin embeddings — determinista y testeable. El upgrade a
recuperación semántica (embeddings) es futuro; ver scripts/mundana/doctrinal_rag.py
en ai-oracle para esa línea.

Formatos soportados: .md / .txt (texto plano), .ts (extrae LILLY_SYSTEM_PROMPT),
.docx (parseado con python-docx si está instalado).
"""
from __future__ import annotations

import re
from dataclasses import dataclass
from functools import lru_cache
from pathlib import Path

import config

# Fuentes doctrinales, relativas a ABU_DOCTRINE_ROOT (repo ai-oracle).
# Cada entrada es un archivo concreto o un glob de directorio. Se escanea el
# corpus doctrinal/conceptual — deliberadamente NO se incluye obsidian_vault
# entero (experimentos, finops, resultados no son doctrina).
_SOURCE_GLOBS: list[str] = [
    # Núcleo
    "AXIOMATICS_OF_HEAVENS_v0.4.md",
    "LILLY_QUOTES.md",
    "next_app/lib/lilly-prompt.ts",
    # Axiomática y mecanismo
    "obsidian_vault/02_doctrina/*.md",
    "obsidian_vault/04_hipotesis/*.md",
    "obsidian_vault/persian_techniques.md",
    # Conceptos + Canon (incluye .docx)
    "docs/concepts/*.md",
    "docs/concepts/*.docx",
    # Teoría / formalización / diálogos de paradigma
    "docs/theory/*.md",
]

_STOPWORDS = {
    "the", "and", "for", "que", "con", "los", "las", "del", "una", "uno",
    "por", "como", "más", "este", "esta", "the", "a", "de", "en", "el",
    "la", "un", "is", "of", "to", "in", "su", "se", "al", "lo",
}


@dataclass
class Fragment:
    fuente: str
    titulo: str
    texto: str


def _tokenize(text: str) -> list[str]:
    words = re.findall(r"[a-záéíóúñü]+", text.lower())
    return [w for w in words if len(w) > 2 and w not in _STOPWORDS]


# ── Lectores por formato ────────────────────────────────────────────────────

def _read_text(path: Path) -> str:
    return path.read_text(encoding="utf-8", errors="replace")


def _read_docx(path: Path) -> str:
    """Extrae texto de un .docx. Devuelve '' si python-docx no está instalado."""
    try:
        from docx import Document  # python-docx
    except ImportError:
        return ""
    try:
        doc = Document(str(path))
    except Exception:
        return ""
    parts = [p.text for p in doc.paragraphs if p.text.strip()]
    for table in doc.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells if c.text.strip()]
            if cells:
                parts.append(" · ".join(cells))
    return "\n".join(parts)


def _read_source(path: Path) -> str:
    if path.suffix == ".docx":
        return _read_docx(path)
    return _read_text(path)


# ── Particionado en secciones ───────────────────────────────────────────────

def _split_ts_prompt(content: str) -> list[tuple[str, str]]:
    """Extrae el string LILLY_SYSTEM_PROMPT del .ts y lo parte por secciones."""
    m = re.search(r"LILLY_SYSTEM_PROMPT\s*=\s*`(.*?)`", content, re.DOTALL)
    body = m.group(1) if m else content
    return _split_blocks(body)


def _split_blocks(content: str) -> list[tuple[str, str]]:
    """Parte texto en (titulo, bloque) por headings markdown o líneas en MAYÚSCULAS."""
    lines = content.splitlines()
    blocks: list[tuple[str, str]] = []
    title = "(intro)"
    buf: list[str] = []

    def flush() -> None:
        text = "\n".join(buf).strip()
        if text:
            blocks.append((title, text))

    for line in lines:
        stripped = line.strip()
        is_md_heading = stripped.startswith("#")
        is_caps_heading = bool(
            stripped
            and len(stripped) < 60
            and re.match(r"^[0-9]*\.?\s*[A-ZÁÉÍÓÚÑ][A-ZÁÉÍÓÚÑ /—-]+$", stripped)
        )
        if is_md_heading or is_caps_heading:
            flush()
            title = stripped.lstrip("#").strip() or title
            buf = []
        else:
            buf.append(line)
    flush()
    return blocks


def _chunk(text: str, max_chars: int) -> list[str]:
    if len(text) <= max_chars:
        return [text]
    parts, cur = [], ""
    for para in text.split("\n\n"):
        if len(cur) + len(para) + 2 > max_chars and cur:
            parts.append(cur.strip())
            cur = para
        else:
            cur = f"{cur}\n\n{para}" if cur else para
    if cur.strip():
        parts.append(cur.strip())
    return parts


def _resolve_sources() -> list[Path]:
    """Expande los globs a rutas concretas existentes, deduplicadas y ordenadas."""
    root = config.ABU_DOCTRINE_ROOT
    seen: dict[str, Path] = {}
    for pattern in _SOURCE_GLOBS:
        if any(ch in pattern for ch in "*?["):
            matches = root.glob(pattern)
        else:
            matches = [root / pattern]
        for p in matches:
            if p.exists() and p.is_file():
                seen[str(p.resolve())] = p
    return sorted(seen.values(), key=lambda p: str(p))


@lru_cache(maxsize=1)
def _index() -> list[Fragment]:
    fragments: list[Fragment] = []
    root: Path = config.ABU_DOCTRINE_ROOT
    for path in _resolve_sources():
        content = _read_source(path)
        if not content.strip():
            continue
        rel = str(path.relative_to(root)) if root in path.parents else path.name
        blocks = _split_ts_prompt(content) if path.suffix == ".ts" else _split_blocks(content)
        for title, text in blocks:
            for chunk in _chunk(text, max_chars=1200):
                fragments.append(Fragment(fuente=rel, titulo=title, texto=chunk))
    return fragments


def retrieve(tema: str, tags: list[str] | None = None, k: int = 3) -> list[dict]:
    """
    Recupera los k fragmentos doctrinales más relevantes para tema + tags.

    Returns: [{fuente, titulo, fragmento, score}] ordenado por score desc.
    """
    query_terms = _tokenize(tema) + [t.lower() for t in (tags or [])]
    if not query_terms:
        return []
    query_set = set(query_terms)

    scored: list[tuple[float, Fragment]] = []
    for frag in _index():
        hay_tokens = set(_tokenize(frag.titulo + " " + frag.texto))
        overlap = query_set & hay_tokens
        if not overlap:
            continue
        score = float(len(overlap))
        title_tokens = set(_tokenize(frag.titulo))
        score += 0.5 * len(query_set & title_tokens)
        scored.append((score, frag))

    scored.sort(key=lambda x: x[0], reverse=True)
    return [
        {
            "fuente": frag.fuente,
            "titulo": frag.titulo,
            "fragmento": frag.texto,
            "score": round(score, 2),
        }
        for score, frag in scored[: max(1, k)]
    ]


def source_status() -> list[dict]:
    """Diagnóstico: qué fuentes doctrinales se resolvieron en disco."""
    root = config.ABU_DOCTRINE_ROOT
    out = []
    for path in _resolve_sources():
        rel = str(path.relative_to(root)) if root in path.parents else path.name
        out.append({"fuente": rel, "encontrada": True, "formato": path.suffix})
    return out
